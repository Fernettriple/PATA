#!/usr/bin/env python3

import docx,os,pyperclip,re,shelve,math,csv,xlsxwriter,openpyxl,shutil
import pandas as pd


#Gloriosas funciones
def purge(x):#CAMBIAR PARA QUE ACEPTE PARAMETROS
    sacar=[' ',',']
    x=x.upper()
    for i in range(len(x)):
        for i in range(len(sacar)):
            x=x.strip(sacar[i])
    return x
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return ''.join(fullText)
def mostrarlista():
    barra()
    for i in range(len(COOP)):
        if i < 10:
            print(str(i)+'. '+TITLE[i].ljust(30,'.')+COOP[i].rjust(40, '.'))
        else:
            print(str(i)+'. '+TITLE[i].ljust(29,'.')+COOP[i].rjust(40, '.'))
    barra()
def purge_nums(a,b,c):
    a=str(a)
    a=a.split(b)
    a=a[1]
    a=a.split(c)
    a=a[0]
    for _ in range(len(a)):
        a=a.strip()
    return a
def getData(a,b,c):
    try:
        a=a.split(b)
        a=a[1]
    except:
        print('No se encontro en el acta la frase clave '+str(b))
        print('Esto es un error CRITICO que es imposible de reparar sin modificar el codigo. Por favor retire la cooperacion actual de la carpeta y hagala a mano(COMO DIOS MANDA)')
    try:
        a=a.split(c)
        a=a[0]
        a=purge(a)
    except:
        print('No se encontro en el acta la frase clave '+str(c))
        print('Esto es un error CRITICO que es imposible de reparar sin modificar el codigo. Por favor retire la cooperacion actual de la carpeta y hagala a mano(COMO DIOS MANDA)')
    return a

def barra():
    print('~'*80)
#-------------------------------------------------------------------------------------------------------------------------------

#Datos necesarios
Nums= {'DI BLASI':'29076875',
       'SALBATIERRA': '35096157',
       'FERNANDEZ': '35417114',
       'REHAK': '31453',
       'ZORNETTA' : '35246371',
       'RODRIGUEZ': '4032',
       'ALARCON' : '34270825',
       'LOPEZ':'26.380.362'}
Drogas=['CLORHIDRATO',
        'BASE', 
        'MARIHUANA',
        'MDMA', 
        'KETAMINA', 
        'METANFETAMINA']
Meses={'ENERO':'01',
       'FEBRERO':'02',
       'MARZO':'03',
       'ABRIL':'04',
       'MAYO' :'05',
       'JUNIO':'06',
       'JULIO':'07',
       'AGOSTO':'08',
       'SEPTIEMBRE':'09',
       'OCTUBRE':'10',
       'NOVIEMBRE':'11',
       'DICIEMBRE':'12'}
Cantidad_de_cooperaciones_leidas=[]
#-------------------------------------------------------------------------------------------------------------------------------

#Creacion de los DF con sus respetivos titulos (Los titulos son usados despues, por eso no los puse adentro)
Titulo_PROD=    ['Nro',
                'Dependencia',
                'Srio. Nro.',
                'Caratula',
                'Cria/Comuna',
                'Magistrado Interventor',
                'Lugar',
                'Peritos',
                'Labor Realizada']
PROD = pd.DataFrame(index=Titulo_PROD)

Titulo_ODI92=[  'N° de procedimiento',
                'Fecha',
                'Hora',
                'JUZGADO/FISCALIA',
                'SECRETARIA',
                'DEPENDENCIA',
                'SUMARIO/CAUSA',
                'Estupefacientes principales',
                'Otros estupefacientes incautados',
                'Cantidad',
                'Unidad de medida']
ODI92=pd.DataFrame(index=Titulo_ODI92)

Titulo_LIBRO=[  'COOPERACION',
                'FECHA',
                'PERITO',
                'JUZGADO/FISCALIA',
                'JUEZ/FISCAL',
                'SECRETARIA',
                'SECRETARIO',
                'NUMERO DE SUMARIO',
                'CARATURA',
                'DAMNIFICADO',
                'IMPUTADO',
                'MATERIAL RECIBIDO',
                'PERITO',
                'DEPENDENCIA QUE RECIBE EL MATERIAL']
LIBRO=pd.DataFrame(index=Titulo_LIBRO)
#-------------------------------------------------------------------------------------------------------------------------------

#MENSAJE DE BIENVENIDA
#Mensaje de bienvenida
print('''Bienvenido al PATA (Programa de Automatizacion de Tareas Administrativas) version 1.0
Este programa sirve para leer Actas en formado .DOCX y crear automaticamente las hojas de ruta, junto con un Excel llamado 'INFO' que en las 3 paginas te deja listo la informacion para pasar a la ODI92, Productividad y Libro virtual
ACLARACION: solo funciona para archivos que se encuentren en EL MISMO DIRECTORIO que este programa
Cada vez que haya un error(puede ser por error del programa, o porque los oficiales escriben CUALQUIER COSA),
el programa mostrara la seccion donde DEBERIA estar la informacion y te va a pedir que la introduzcas manualmente.
Por favor, ante cualquier duda, consulta''')
barra()
print('Desea correr el programa en "Modo Seguro"? (Modo seguro te muestra todos los datos leidos ANTES de pasarlos a los excels')
print('Si/No')
barra()
ñ=input()
ñ=ñ.upper()
U=0
#-------------------------------------------------------------------------------------------------------------------------------

for filename in os.listdir('.'):
    if filename=='HDR.docx':
        continue
    elif filename.startswith("HOJA"):
        continue
    elif filename.endswith(".docx"):        
        Title= ['Cooperacion',
                'Sumario',
                'LP',
                'Suscribiente', 
                'Comiseria/Division/etc',
                'Fiscalia/Juzgado',
                'Magistrado interventor',
                'Secretaria', 
                'Secretario',
                'Imputado/s',
                'Perito',   
                'Fecha',
                'Hora de inicio',
                'Hora de finalizacion']
        TITLE=['Cooperacion',
                'Sumario',
                'LP',
                'Suscribiente', 
                'Comiseria/Division/etc',
                'Fiscalia/Juzgado',
                'Magistrado interventor',
                'Secretaria', 
                'Secretario',
                'Imputado/s',
                'Perito',   
                'Fecha',
                'Hora de inicio',
                'Hora de finalizacion',
                'Peso de MDMA',
                'Peso de Marihuana',
                'Peso de Clorhidrato de Cocaina',
                'Peso de Base de Cocaina']
        while True:
            try:
                doc=getText(filename)
                doc=doc.upper()
                docfull= docx.Document(filename)
                break
            except:
                print('Hubo un error inesperado (ni idea que puede ser) o el archivo que ingresaste NO ES un archivo .docx')
                print('Pruebe de nuevo. Si no funciona y vuelve a ver este mensaje es que se rompio el programa o hay algo muy raro con el archivo que esta leyendo('+filename+').'+
                '\n Por favor saquelo de la carpeta y copielo manualmente(vuelva a iniciar el programa)')
                break
        #-------------------------------------------------------------------------------------------------------------------------------

        #Numero de Cooperacion y Sumario

        try:
            section = docfull.sections[0]
            header = section.header
            spam=header.paragraphs
            bacon=0
            if bacon<2:
                for paragraph in header.paragraphs:
                    if 'COOPERACION' in paragraph.text:
                        spam=paragraph.text
                    if 'SUMARIO' in paragraph.text:
                        eggs=paragraph.text
                        bacon+=1

            spam=purge_nums(spam,':','/')
            sumario=re.compile(r'(\d)*/((\d)*)?')
            _=sumario.search(eggs)
            COOP=[]
            NumCoop= purge(spam)
            COOP.append(NumCoop)
            COOP.append(_.group())
        except:
            print('----'*5)
            print('Error en la lectura del encabezado. Por favor, introduzcalo manualmente')
            section = docfull.sections[0]
            header = section.header
            spam=header.paragraphs
            print('Aca esta el fragmento donde DEBERIA estar los numeros necesarios\n'+spam)
            print('Numero de cooperacion:')
            spam=input()
            print('Numero de sumario:')
            eggs=input()
            COOP.append(purge(spam))
            COOP.append(purge(eggs))
        barra()
        print('ACTUALMENTE LEYENDO LA COOPERACION : '+spam)
        barra()
        #-------------------------------------------------------------------------------------------------------------------------------

        #LP
        try:
            spam=getData(doc,'SUSCRIBE','DEL NUMERARIO')
            numLP= re.compile(r'(\d)+')
            bugg=re.compile(r'°1|1°|º1|1º|(\s)1(\s)')
            if bugg.search(spam):
                bug=bugg.search(spam)
                spam=spam.split()
                spam.remove(bug.group())
                spam=' '.join(spam)
            LPS=numLP.search(spam).group()
            COOP.append(LPS)
                     
        except:
            print('Error al leer el LP del suscribiente. Por favor, introduzcalo manualmente')
            spam=getData(doc,'EL FUNCIONARIO','A LOS FINES LEGALES')
            print('Aca esta el fragmento donde DEBERIA estar los numeros necesarios\n'+spam)
            print('LP:')
            LPS=input()
            COOP.append(LPS)
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #Suscribiente    
        try:    
            spam=getData(doc,'SUSCRIBE','DEL NUMERARIO')
            eggs=re.compile(r'L(.)?P(.)?[^RIMERO]?')
            spam=spam.split()
            for i in range(len(spam)):
                try:
                    if eggs.search(spam[i]):
                        _=eggs.search(spam[i])
                        _=_.group()
                        _=_.strip(' ')
                        spam.remove(_)
                except:
                    continue
            spam.remove(LPS)
            spam=' '.join(spam)
            COOP.append(spam)
        except:
            print('Error a leer el nombre de suscribiente. Por favor introduzcalo manualmente')
            spam=getData(doc,'EL FUNCIONARIO','A LOS FINES LEGALES')
            print('Aca esta el fragmento donde DEBERIA estar los numeros necesarios\n'+spam)
            spam=input()
            spam=spam.upper()
            COOP.append(spam)

        #-------------------------------------------------------------------------------------------------------------------------------

        #Comiseria/Division/etc
        try:
            spam=getData(doc,'NUMERARIO DE','A LOS FINES')
            NumCC=re.compile(r'((\d)?\d)(\w)?')
            CC=NumCC.search(spam).group()
            CC='CC'+CC            
            COOP.append(purge(spam))
        except:
            print('Error al leer el nombre de la comiseria/division/etc. Por favor, introduzcalo manualmente:')
            spam=getData(doc,'EL FUNCIONARIO','A LOS FINES LEGALES')
            print('Aca esta el fragmento donde DEBERIA estar los numeros necesarios\n'+spam)
            spam=input()
            spam=spam.upper()
            if NumCC.search(spam):
                CC=NumCC.search(spam).group()
                CC='CC'+CC
            else:
                CC=spam            
            COOP.append(purge(spam))
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #Magisterio interventor
        try:
            MagInterventor=getData(doc,'CON INTERVENCIÓN DE','SECRE')
            SecrInterventora=getData(doc,'SECRE','EN LA QUE RES')
            fiscalia=re.compile(r'(FISCAL.A)')
            juzgado=re.compile(r'(JUZGADO)')
            NumMagisterio=re.compile(r'((\d)?\d)')
            dr=re.compile(r'(DR(.*)?(\w)+)')
            if fiscalia.search(MagInterventor):
                M='F'       #Tipo de Magisterio. Si es Fiscalia o Juzgado (me sirve para dps)          
                _=NumMagisterio.search(MagInterventor).group()
                NumFiscalia='F'+_
                COOP.append(_)
                if dr.search(MagInterventor):
                    _=dr.search(MagInterventor).group()
                    if ',' in _:
                        _=_.split(',')
                        _=_[0]
                    COOP.append(_)
                else:
                    print('Error:No se ha encontrado el nombre del fiscal. Introducilo Manualmente')
                    print('Aca esta el fragmento donde DEBERIA estar el nombre del fiscal\n'+MagInterventor)
                    _=input()
                    COOP.append(_)
                COOP.append('UNICA')
                if dr.search(SecrInterventora):
                    _=dr.search(SecrInterventora).group()
                    if ',' in _:
                        _=_.split(',')
                        _=_[0]
                    COOP.append(_)
                else:
                    print('Error:No se ha encontrado el nombre del secretario. Introducilo Manualmente')
                    print('Aca esta el fragmento donde DEBERIA estar el nombre del secretario\n'+SecrInterventora)
                    _=input()
                    COOP.append(_) 
            elif juzgado.search(MagInterventor):
                M='J' #Tipo de Magisterio. Si es Fiscalia o Juzgado (me sirve para dps)
                _=NumMagisterio.search(MagInterventor).group()
                COOP.append(_)
                NumJuzgado='J'+_
                if dr.search(MagInterventor):
                    _=dr.search(MagInterventor).group()
                    if ',' in _:
                        _=_.split(',')
                        _=_[0]
                    COOP.append(_)
                else:
                    print('Error:No se ha encontrado el nombre del juez. Introducilo manualmente ')
                    print('Aca esta el fragmento donde DEBERIA estar el nombre del juez\n'+MagInterventor)
                    _=input()
                    COOP.append(_)
                _=NumMagisterio.search(SecrInterventora).group()
                COOP.append(_)
                NumJuzgadoySecr=NumJuzgado+' S'+_
                if dr.search(SecrInterventora):
                    _=dr.search(SecrInterventora).group()
                    if ',' in _:
                        _=_.split(',')
                        _=_[0]
                    COOP.append(_)
                else:
                    print('Error:No se ha encontrado el nombre del secretario. Introducilo Manualmente')
                    print('Aca esta el fragmento donde DEBERIA estar el nombre del secretario\n'+SecrInterventora)
                    _=input()
                    COOP.append(_) 
        except:
            print('Algo terrible paso con la parte del magisterio interventor. Rellenalo manualmente')
            print('Primero introduzca J o F si es Juzgado o Fiscalia. Si es cualquier otra cosa, este programa no lo soporta. Saque esa Cooperacion y vuelva a correr este programa')
            M=input()
            print('Numero de Fiscalia/Juzgado?')
            COOP.append(input())
            print('Nombre del Fiscal/Juez?')
            COOP.append(input())
            print('Numero de Secretaria?(en caso de ser fiscalia poner "UNICA"')
            COOP.append(input())
            print('Nombre del secretario?')
            COOP.append(input())
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #Imputado VER COMO RESOLVER ESTO.
        try:
            print('Por ahora no se como distinguir el imputado.\n'+
                'Por ahora necesito que lo escribas manualmente.\n'+
                'Para ayudarte, aca tenes el fragmento del acta que habla de los imputados. Gracias y disculpas!')
            spam=getData(doc,'IMPUTAD','A FIN DE')
            print(spam)
            eggs=input()
            eggs=purge(eggs)
            COOP.append(eggs)
        except:
            print('''Error al introducir imputados. Mira el acta e introducilo manualmente porque ocurrio un problema bastante grave a la hora de leer el acta''')
            COOP.append(input())
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #Peritos
        _=0
        try:
            spam=getData(doc,'AGREGANDO QUE PARA','SE UTILIZARÁN LOS SIGUIENTES REACTIVOS')
            for k in Nums.keys():
                if k in spam:
                    COOP.append(k)
                    _+=1
            if _==0:
                print('Ocurrio un error a la hora de leer el nombre del perito, por favor introduzcalo ')
                COOP.append(input())
        except:
            print('Error CATASTROFICO a la hora de leer el acta en la parte del perito. Debe haber algun error importante en el modelo. Por favor, introduzca el perito manualmente')
            COOP.append(input())
        #-------------------------------------------------------------------------------------------------------------------------------    
                      
        #Fecha, Hora de inicio y finalizacion
        ##Fecha
        try:
            spam=getData(doc,'AIRES, HOY',' SIENDO LAS')
            spam=spam.split()
            dia=spam[0]
            for k,v in Meses.items():
                if k == spam[4]:
                    mes = v
            año=spam[7]
            _=dia+'/'+mes+'/'+año
            COOP.append(_)
        except:
            print('''Ha habido un error debido a que la fecha se encuentra mal escrita
        Para que la lea el programa, expresarla de la siguiente manera:
        hoy XX del Mes de XXXX del año XXXX
        Igual, te la dejo pasar por esta vez, introduci la fecha manualmente en formato digital
        Por ejemplo: 09/07/1816
        Aca te dejo el fragmento del acta donde deberia esta la informacion''')
            eggs=getData(doc,'AIRES, HOY',' SIENDO LAS')
            print(eggs)
            spam=input()
            COOP.append(spam)

        ##Hora inicial y final
        try:
            spam=getData(doc,'SIENDO LAS','HORAS')
            if len(spam)==5:
                spam=spam.replace(',',':')
                spam=spam.replace('.',':')
                COOP.append(spam)
            else:
                print('''Error en el formato de la hora inicial. El programa solo admite XX:XX
            Introduzca la hora inicial''')
                COOP.append(input())
            spam=getData(doc,'TERMINADO EL ACTO, SIENDO LAS','HORAS')
            if len(spam)==5:
                spam=spam.replace(',',':')
                spam=spam.replace('.',':')
                COOP.append(spam)
            else:
                print('''Error en el formato de la hora final. El programa solo admite XX:XX 
                Introduzca la hora de finalizacion''')
                COOP.append(input())
        except:
            print('Error al leer la hora de inicio o finalizacion en el acta.\n'+
                'Por favor, introduzcalas manualmente en el formato XX:XX')
            print('Hora inicial:')
            COOP.append(input())
            print('Hora finalizacion:')
            COOP.append(input())
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #Peso de las Drogas. 
        try:
            spam=getData(doc,'ACTO SEGUIDO SE PROCEDE A REALIZAR LA APERTURA DE','FINALIZADO EL PROCEDIMIENTO')
            pesos=re.compile(r'\d*,?\.?\d{3}')
            if 'GRAMOS' in spam:
                bacon=spam.replace('GRAMOS','GRAMOS\n')
            elif 'GRS' in spam:
                bacon=spam.replace('GRS','GRAMOS\n')
            if 'SUSTANCIA' in bacon:               
                bacon=bacon.replace('SUSTANCIA','\nSUSTANCIA')
            if 'MATERIAL' in bacon:
                bacon=bacon.replace('MATERIAL','\nMATERIAL')
            
            bacon=bacon.split('\n') #CON ESTO LOGRO SEPARAR EN FRASES QUE EMPIECEN CON 'SUSTANCIA/MATERIAL' Y TERMINEN CON 'GRAMOS'. SUPONGO QUE ESTO CONTENDRA EL PESO DE LA DROGA
            drg= {}
            if 'MDMA' in doc:
                print('Hay MDMA, tenes que introducir manualmente el peso')
                _=input()
                COOP.append(_)
                Title.append('Peso de MDMA') 
                drg['MDMA']=_    
            else:
                COOP.append('NO HAY')
                Title.append('Peso de MDMA')
            
            Mari=0
            Coca=0
            Paco=0
            for i in range(0,len(bacon)-1):
                if 'SUSTANCIA' in bacon[i] or 'MATERIAL' in bacon[i]:
                    if 'PESA' in bacon[i]:
                        _=pesos.findall(bacon[i])
                        _=_[0]
                        _=str(_)
                        _=_.replace(',','.')
                        _=float(_)
                        mariajuana={'VEGETAL','VERDE AMARRONADA','VERDE','VERDEAMARRONADA','VERDE MARRON'}
                        for v in mariajuana:        
                            if v in bacon[i]:
                                print(bacon[i])
                                Mari += _
                                Mari=round(Mari,3)
                                bacon[i]='wea'
                        if 'BLANCA' in bacon[i]:
                            Coca=Coca+_
                            Coca=round(Coca,3)
                            bacon[i]='wea'
                        if 'AMARI' in bacon[i]:
                            Paco=Paco+_
                            Paco=round(Paco,3)
                            bacon[i]='wea'
                    else:
                        continue
            if Mari != 0:
                Mari=str(Mari)
                Mari=Mari.replace('.',',')
                Title.append('Peso de Marihuana')
                COOP.append(Mari)
                drg['MARIHUANA']=Mari
            else:
                COOP.append('NO HAY')
                Title.append('Peso de Marihuana')            
            if Coca != 0:                
                Coca=str(Coca)
                Coca=Coca.replace('.',',')
                COOP.append(Coca)
                Title.append('Peso de Cocaina')
                drg['CLORHIDRATO DE COCAINA']=Coca 
            else:
                COOP.append('NO HAY')
                Title.append('Peso de Clorhidrato de Cocaina')   
            if Paco !=0:
                Paco=str(Paco)
                Paco=Paco.replace('.',',')
                COOP.append(Paco)
                Title.append('Peso de Base de Cocaina')
                drg['BASE DE COCAINA']=Paco    
            else:
                COOP.append('NO HAY')
                Title.append('Peso de Base de Cocaina')
        except:
            print('Error NEFASTO a la hora de leer las drogas. Por favor, saque la cooperacion y vuelva a correr el programa(algun dia por ahi introduzco el codigo necesario para arreglar este error, si me aumentan el sueldo..)')
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #CHECKEO QUE ESTE TODO BIEN. RE HACER TODO. VER COMO HAGO PARA QUE NO 
        if len(COOP)!=18:
             while True:
                mostrarlista()
                print()
                print('Desea Modificar algo? Si es asi, introduzca el indice (Numero a la izquierda). En el caso de que este todo bien, apretar enter')
                spam=input()
                if spam != '':
                    spam=int(spam)
                    print('Que valor deberia tener esto?')
                    COOP[spam]=input()
                    print()
                else:
                    break
        for i in range(len(COOP)):
            if len(COOP[i])>50:
                print('El programa detecto que el siguiente dato es muy largo, se puede deber a un problema de lectura o por ahi esta bien. Reviselo por las dudas:')
                print(COOP[i])
                print('Si esta todo bien, presione enter. Si desea cambiarlo, presione cualquier otro boton')
                if input()!='':
                    print('Introduzca por favor el valor correcto (o una version mas resumida, sin omitir los datos IMPORTANTES)')
                    COOP[i]=input()
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #Checkeo al final SI ESTA EN "MODO SEGURO"
        if ñ=='SI' or ñ=='S':
            while True:
                mostrarlista()
                print()
                print('Desea Modificar algo? Si es asi, introduzca el indice (Numero a la izquierda). En el caso de que este todo bien, apretar enter')
                spam=input()
                if spam != '':
                    spam=int(spam)
                    print('Que valor deberia tener esto?')
                    COOP[spam]=input()
                    print()
                else:
                    break
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #CREACION DE LISTAS Y DATAFRAMES Y OTRAS COSAS QUE NECESITE HACER PARA QUE FUNCIONE ESTA WEA QLA
        if M=='F': #SI ES FISCALIA, VA F+ EL NUMERO SIN SECRETARIA
            P=[COOP[0],'DIVISION QUIMICA INDUSTRIAL Y ANALISIS FISICOS Y QUIMICOS', COOP[1], 'INF. LEY 23737',CC,NumFiscalia,'EDIFICIO CENTRAL',COOP[10],'TEST ORIENTATIVO Y PESAJE' ]
            O=[COOP[0],COOP[11],COOP[12],NumFiscalia,'S1',CC,COOP[1],]
            L=[COOP[0],COOP[11],COOP[10],NumFiscalia,COOP[6],COOP[7], COOP[8], COOP[1],'INFRACCION A LA LEY 23.737','LEY Y SOCIEDAD',COOP[9],'UN SOBRE', COOP[10],CC]
        elif M=='J': #SI ES JUZGADO, VA J+NUMERO+S+NUMERO
            P=[COOP[0],'DIVISION QUIMICA INDUSTRIAL Y ANALISIS FISICOS Y QUIMICOS', COOP[1], 'INF. LEY 23737',CC,NumJuzgadoySecr,'EDIFICIO CENTRAL',COOP[10],'TEST ORIENTATIVO Y PESAJE' ]
            O=[COOP[0],COOP[11],COOP[12],NumJuzgado,COOP[7],CC,COOP[1]]
            L=[COOP[0],COOP[11],COOP[10],NumJuzgado,COOP[6],COOP[7], COOP[8], COOP[1],'INFRACCION A LA LEY 23.737','LEY Y SOCIEDAD',COOP[9],'UN SOBRE', COOP[10],CC]
        OO_Df= pd.DataFrame(data=None, index=Titulo_ODI92)
        R=U
        i=0
        for k,v in drg.items(): #con esto agrego las drogas encontradas
            if i==0:
                O.append(k)
                O.append('')
                O.append(v)
                O.append('GRAMOS')
                i+=1                
            else:
                R+=1
                U=R
                PP=['','','','','','','',k,'',v,'GRAMOS']
                PP_Df=pd.DataFrame(data=PP,index=Titulo_ODI92)
                OO_Df=OO_Df.join(other=PP_Df, lsuffix='_left', rsuffix='_right')
                OO_Df=OO_Df.dropna()
                PP=[]
                PP_Df={}
        P_Df = pd.DataFrame(data=P,index=Titulo_PROD)
        O_Df= pd.DataFrame(data=O,index=Titulo_ODI92)
        O_Df=O_Df.join(other=OO_Df, lsuffix='_left', rsuffix='_right')
        O_Df=O_Df.dropna()
        L_Df=pd.DataFrame(data=L,index=Titulo_LIBRO)
        ODI92=ODI92.join(other=O_Df, lsuffix='_left', rsuffix='_right')
        LIBRO=LIBRO.join(other=L_Df, lsuffix='_left', rsuffix='_right')
        LIBRO=LIBRO.drop_duplicates() #no se xq me tira muchas veces el perito
        LIBRO=LIBRO.loc[Titulo_LIBRO, :]
        PROD=PROD.join(other=P_Df,lsuffix='_left', rsuffix='_right')
        #-------------------------------------------------------------------------------------------------------------------------------
        
        #ACA VA LA HOJA DE RUTA
        shutil.copy('HDR.docx', 'HOJA DE RUTA DE '+COOP[0]+'.docx')
        hdr= docx.Document('HOJA DE RUTA DE '+COOP[0]+'.docx')
        def replace_string(filename,viejo,nuevo):
            for table in filename.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:                                
                            if viejo in p.text:
                                inline = p.runs
                                # Loop added to work with runs (strings with same style)
                                for i in range(len(inline)):
                                    if viejo in inline[i].text:
                                        text = inline[i].text.replace(viejo, nuevo)
                                        inline[i].text = text
            return 1
                
        
        replace_string(hdr,'!',COOP[0])
        replace_string(hdr,'$',COOP[11])
        replace_string(hdr,'%',COOP[1])
        replace_string(hdr,'&',CC)
        if M=='F':
            replace_string(hdr,'*',('FISCALIA PENAL CONTRAVENCIONAL Y DE FALTAS N°'+COOP[5]+' A CARGO DE LA '+COOP[6]))
            replace_string(hdr,'ç',(COOP[7]+' A CARGO DE '+COOP[8]))
        elif M=='J':
            replace_string(hdr,'*',('JUZGADO FEDERAL BUSCAR QUE TIENE QUE DECIR N°'+COOP[5]+' A CARGO DE LA '+COOP[6]))
            replace_string(hdr,'ç',(COOP[7]+' A CARGO DE '+COOP[8]))
        else:
            print('Error importante, no puedo identificar si es Fiscalia o Juzgado, se dejara esa parte de la hoja de ruta en blanco')
            continue
        replace_string(hdr,'¨',COOP[9])
        replace_string(hdr,'<',(COOP[10]+'  DNI: '+ Nums[COOP[10]]))
        replace_string(hdr,'>',(COOP[3]+' LP: '+COOP[2]))
        replace_string(hdr,'+',COOP[12])
        replace_string(hdr,'-',COOP[13])
        if COOP[14]!='NO HAY':
            replace_string(hdr,'DR1',('MDMA: '+COOP[14]))
        else:
            replace_string(hdr,'DR1','')
        if COOP[15]!='NO HAY':
            replace_string(hdr,'DR2',('M: '+COOP[15]))
        else:
            replace_string(hdr,'DR2','')
        if COOP[16]!='NO HAY':
            replace_string(hdr,'DR3',('CC: '+COOP[16]))
        else:
            replace_string(hdr,'DR3','')
        if COOP[17]!='NO HAY':
            replace_string(hdr,'DR4',('BC: '+COOP[17]))
        else:
            replace_string(hdr,'DR4','')
        hdr.save('HOJA DE RUTA DE '+COOP[0]+'.docx')                
        U+=1
        L=[]
        O=[]
        P=[]
        PP=[]
        Cantidad_de_cooperaciones_leidas=Cantidad_de_cooperaciones_leidas.append(COOP[0])
        continue
    else:
        continue
        
#-------------------------------------------------------------------------------------------------------------------------------
#ACA AGARRO LOS DF Y LOS PASO AL EXCEL   
PROD=PROD.dropna()
PROD=PROD.transpose()
ODI92=ODI92.dropna()
ODI92=ODI92.transpose()
LIBRO=LIBRO.dropna()
LIBRO=LIBRO.transpose()
with pd.ExcelWriter('Info.xlsx') as writer:  
    PROD.to_excel(writer, sheet_name='PROD')
    ODI92.to_excel(writer, sheet_name='ODI92')
    LIBRO.to_excel(writer, sheet_name='LIBRO')
barra()
print('Programa terminado. :^D')
Cantidad_de_cooperaciones_leidas=', '.join(Cantidad_de_cooperaciones_leidas)
print('Se leyeron las siguientes cooperaciones: '+Cantidad_de_cooperaciones_leidas)
barra()
##TO DO LIST:
## REEMPLAZAR TODOS LOS .APPEND() POR DEFINIR UNA LISTA AL PRINCIPIO CON TODOS LOS VALORES COMO
#  "ERROR DE TAL COSA" Y EN LUAR DE IR APPENDEANDO, IR REEMPLAZANDO. ES MEJOR PORQ ME EVITA
#  ERRORES DE LONGITUD DE LISTA Y ES MAS FACIL RASTREAR LOS ERRORES Y REEMPLAZARLOS

